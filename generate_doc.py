"""Generate project summary Word document for 自动视频剪辑工具."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style helpers ──
style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    # Clear default run and add our own
    p.clear()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    return p

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading("自动视频剪辑工具 — 项目总结文档", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"

para("")
para("基于 AI 多轮对话的视频自动剪辑平台", bold=True)
para("版本：v1.0  |  日期：2026 年 5 月 9 日")
para("技术栈：faster-whisper + pyannote + DeepSeek + FFmpeg + FastAPI")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 软件概述
# ═══════════════════════════════════════════════════════════════
heading("一、软件概述", level=1)

para("本软件是一套基于 AI 语音识别与自然语言理解的视频自动剪辑系统。用户只需上传视频，系统会自动完成语音转写、说话人识别，然后用户可以用日常大白话（如「把讲产品优势的部分留下来」）描述剪辑需求，AI 自动筛选出相关片段并支持导出。")

para("核心思路：用 AI 替代传统视频剪辑软件中繁琐的手动时间轴拖拽操作，把「剪辑」变成「对话」。")

# ═══════════════════════════════════════════════════════════════
# 2. 最终实现的功能
# ═══════════════════════════════════════════════════════════════
heading("二、已实现的最终功能", level=1)

features = [
    ("多语言语音转写",
     "基于 faster-whisper（medium 模型），支持中、英、日、韩等 99 种语言的自动语音识别，输出带时间戳的分段文字稿。"),
    ("说话人分离",
     "基于 pyannote.audio（speaker-diarization-3.1），自动识别视频中不同说话人，给每段文字打上说话人标签，支持双击改名（如「主持人」「嘉宾A」）。"),
    ("AI 多轮对话剪辑",
     "接入 DeepSeek 大模型，用户在对话框用自然语言描述需求（如「只留下讲产品优势的部分」「去掉开场寒喧」），AI 自动匹配相关文字段落并返回选中的片段 ID。支持多轮对话，上下文连贯，可逐步缩小范围。"),
    ("文字稿搜索与视频跳转",
     "在文字稿区域输入关键词，自动高亮匹配段落，回车跳转视频到对应时间点，支持上下箭头在匹配项之间导航。"),
    ("手动选择与全选",
     "点击任意文字段落即可选中/取消，支持全选、取消全选，选中段落蓝色高亮。同时视频播放器自动跳转到该段落对应时间。"),
    ("拼接导出视频",
     "将选中的视频片段按时间顺序拼接，导出为 MP4 文件。采用帧精确重编码切割（libx264 ultrafast），确保无画面重复或卡顿。"),
    ("带字幕烧录导出",
     "在拼接导出的同时，将文字稿以字幕形式烧录到视频画面上（FFmpeg drawtext），支持字体大小、颜色、描边等样式。"),
    ("SRT 字幕文件导出",
     "导出标准 SRT 格式字幕文件，可直接导入 Premiere、Final Cut Pro、DaVinci Resolve 等专业软件使用。"),
    ("开机自启服务",
     "通过 macOS launchd 配置为开机自启后台服务（KeepAlive），服务器崩溃自动重启，无需手动维护。"),
]

for title_text, desc in features:
    bullet(f"【{title_text}】{desc}")

# ═══════════════════════════════════════════════════════════════
# 3. 解决的问题
# ═══════════════════════════════════════════════════════════════
heading("三、软件解决了什么问题", level=1)

problems = [
    "视频剪辑门槛高：传统剪辑软件（PR、FCPX、达芬奇）学习曲线陡峭，非专业人士需要数周甚至数月才能上手基础操作。本软件将剪辑变为「对话」，零学习成本。",
    "长视频素材筛选费时：一小时的会议录屏、访谈视频，人工回看 + 手动剪辑可能需要 2-3 小时。AI 转写后对话式筛选，可缩短到几分钟。",
    "多语言视频处理困难：市面多数语音识别工具仅支持单一语言或需要手动切换。faster-whisper 自动识别语种，无需预设。",
    "多人对话无法区分：传统转写工具将所有人说话混在一起，无法区分谁说了什么。说话人分离解决了这个问题。",
    "非技术用户的服务部署困难：普通用户不理解终端、Python、服务器等概念。通过 launchd 开机自启，用户只需打开浏览器即可使用。",
]

for p in problems:
    bullet(p)

# ═══════════════════════════════════════════════════════════════
# 4. 业务应用场景
# ═══════════════════════════════════════════════════════════════
heading("四、可应用的业务场景", level=1)

scenarios = [
    ("自媒体/短视频创作",
     "创作者拍摄大量口播素材，通过 AI 对话快速筛选出精彩片段，一键导出。支持多语言转写，适合面向海外平台的创作者。说话人分离功能适合双人/多人访谈类账号。"),
    ("企业内部会议纪要",
     "会议录像上传后可自动生成带时间戳的文字稿 + 说话人标签，快速定位关键讨论片段，导出字幕版视频用于分享。配合 SRT 字幕导出，可进一步整理为会议纪要文档。"),
    ("教育培训",
     "教师录制的课程视频可通过 AI 筛选重点知识点片段，拼接成精华版复习视频。学生可搜索关键词快速跳转到视频中对应知识点位置。"),
    ("市场调研/用户访谈",
     "研究人员录制用户访谈后，AI 转写并区分主持人和受访者的发言，快速定位有价值观点，大幅减少回看时间。对话式剪辑可筛选特定话题相关片段。"),
    ("法律取证/合规审查",
     "长时录音视频通过 AI 转写和搜索快速找到关键对话片段，说话人分离帮助确认发言者身份，SRT 字幕可作为文字证据附件。"),
    ("播客/Podcast 制作",
     "多嘉宾播客录制后的剪辑工作——去掉口误、空白、跑题段落，保留精华内容。AI 对话剪辑比手动拖时间轴高效 10 倍以上。"),
]

for title_text, desc in scenarios:
    bullet(f"【{title_text}】{desc}")

# ═══════════════════════════════════════════════════════════════
# 5. 技术架构
# ═══════════════════════════════════════════════════════════════
heading("五、技术架构", level=1)

para("整体采用「Python 后端 + 单页 HTML 前端」的轻量架构：", bold=False)
bullet("后端：FastAPI（异步 Web 框架）+ uvicorn（ASGI 服务器）")
bullet("语音识别：faster-whisper（本地 Whisper 推理，无需联网）")
bullet("说话人分离：pyannote.audio（speaker-diarization-3.1 模型）")
bullet("AI 语义理解：DeepSeek Chat API（OpenAI 兼容接口）")
bullet("视频处理：FFmpeg（音频提取、视频切割、拼接、字幕烧录）")
bullet("前端：单页 HTML + Tailwind CSS + 原生 JavaScript（无框架依赖）")
bullet("部署：macOS launchd 开机自启服务")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. 开发过程中遇到的卡点与解决方案
# ═══════════════════════════════════════════════════════════════
heading("六、开发过程中的卡点与解决方案", level=1)

challenges = [
    ("卡点 1：Python 3.9 兼容性问题",
     "问题：macOS 自带 Python 3.9.6，不支持 Python 3.10+ 的 X | None 联合类型语法，导致多个文件 import 时报 SyntaxError。",
     "解决：在所有 .py 文件顶部添加 from __future__ import annotations，使得类型注解的求值延迟到运行时，兼容 Python 3.9。"),

    ("卡点 2：NumPy 版本冲突导致崩溃",
     "问题：faster-whisper 基于 NumPy 1.x 编译，但系统安装了 NumPy 2.x，运行时 ABI 不兼容导致进程崩溃退出。",
     "解决：使用 pip3 install 'numpy<2' 将 NumPy 降级到 1.26.4，与 faster-whisper 的预编译 wheel 兼容。"),

    ("卡点 3：FFmpeg 无法通过 Homebrew 安装",
     "问题：用户网络连接 Homebrew 服务器不稳定，多次尝试 brew install ffmpeg 均因 SSL/网络错误失败。",
     "解决：放弃 Homebrew，改用 Python urllib 直接从 evermeet.cx 下载 FFmpeg 静态编译二进制文件，放到 /usr/local/bin/ 目录下。同时在 video_processor.py 中使用 shutil.which() 动态查找 FFmpeg 路径，确保 launchd 环境下也能找到。"),

    ("卡点 4：服务器终端关闭后进程被杀（OpenMP 崩溃）",
     "问题：通过终端 uvicorn 启动服务后，关闭终端窗口时服务器进程被终止。且运行时出现 OMP Error #15（OpenMP 库冲突），导致 CPU 资源初始化失败，服务崩溃。",
     "解决：① 在启动命令中添加环境变量 KMP_DUPLICATE_LIB_OK=TRUE（允许 OpenMP 重复加载）；限制 OMP_NUM_THREADS=1（避免多线程冲突）。② 创建 macOS launchd plist 配置文件，将服务注册为系统后台守护进程，设置 KeepAlive: true 确保崩溃自动重启。"),

    ("卡点 5：launchd 无法读取用户级文件（Operation not permitted）",
     "问题：项目初始放在桌面 ~/Desktop/ 文件夹下，macOS 的 TCC（透明、同意和控制）安全策略禁止 launchd 访问桌面文件。",
     "解决：将整个项目目录移动到 /Users/bdh/auto-video-editor/（用户主目录下的非受保护目录），更新 launchd plist 中的 WorkingDirectory 路径。"),

    ("卡点 6：launchd 环境中找不到 FFmpeg",
     "问题：launchd 守护进程不继承用户的 shell PATH 环境变量，导致 /usr/local/bin/ffmpeg 无法被找到，视频处理全部报错。",
     "解决：① 在 launchd plist 的 EnvironmentVariables 中显式设置 PATH=/usr/local/bin:/usr/bin:/bin。② 在 video_processor.py 中优先使用 shutil.which('ffmpeg') 查找，找不到再回退到硬编码路径。"),

    ("卡点 7：Pyannote 模型下载与 Token 认证",
     "问题：speaker-diarization-3.1 模型托管在 HuggingFace 上，需要：① 接受模型使用条款；② 提供有效的 HF_TOKEN。缺一不可，否则模型加载时返回 'NoneType' object is not callable 错误。且 huggingface_hub 新版本移除了 use_auth_token 参数。",
     "解决：① 在 HuggingFace 网站上分别接受 pyannote/speaker-diarization-3.1 和 pyannote/segmentation-3.0 两个模型页面的使用条款。② 将 huggingface_hub 降级到 0.26.5 版本以兼容 use_auth_token 参数。③ 在 .env 文件和 launchd plist 的 EnvironmentVariables 中都配置 HF_TOKEN。④ 添加 try/except 异常捕获，语音识别失败时给出明确警告而非崩溃。"),

    ("卡点 8：依赖库冲突（k2 / speechbrain / torch）",
     "问题：k2 1.24.1 依赖 torch 1.13.1，但 pyannote.audio 需要 torch >= 2.0，两者不兼容导致安装陷入死循环。",
     "解决：卸载 k2 和 speechbrain（pyannote 实际不依赖它们），只保留 torch >= 2.0，pyannote 正常运行。"),

    ("卡点 9：视频导出画面重复播放",
     "问题：使用 FFmpeg -c copy（流复制）模式切割视频时，只能在关键帧（I 帧）位置切割，导致每个片段开头包含了上一个片段的残余帧，拼接后画面出现重复/回跳。",
     "解决：将切割方式从流复制改为重编码（libx264 + preset ultrafast + crf 23）。同时采用两阶段精确寻址：先用 -ss 快速跳到关键帧附近，再用输出端 -ss 精确裁剪到目标时间点。在保证帧级精确的前提下，ultrafast 预设使重编码速度可控。"),

    ("卡点 10：Whisper 转写速度过慢",
     "问题：Intel i5 CPU 上跑 Whisper medium 模型，原始参数下（beam_size=5, word_timestamps=True, 无 VAD）一个 5 分钟视频转写可能需要 10-20 分钟，用户以为系统卡死。",
     "解决：① 优化推理参数：beam_size=1（贪心解码）、word_timestamps=False（不需要逐词时间戳）、vad_filter=True（跳过静音）、best_of=1。综合提速约 2-4 倍。② 启动时预加载模型到内存（@app.on_event('startup')），避免首次请求时等待模型加载。③ 添加详细的计时日志，方便定位瓶颈。"),

    ("卡点 11：网络不稳定导致模型下载失败",
     "问题：用户的网络环境访问 HuggingFace 不稳定，多次尝试下载小型模型（tiny/base/small）均失败，只有 medium 模型在之前网络正常时下载成功。无法通过下载更小的模型来进一步提速。",
     "解决：① 使用本地已缓存的 medium 模型，通过参数优化弥补速度劣势。② 在 launchd plist 中预先配置 HF_TOKEN，避免运行时认证失败。③ 后续可考虑在国内模型镜像站（如 modelscope）下载模型。"),
]

for title_text, problem_text, solution_text in challenges:
    heading(title_text, level=2)
    para("❌ 问题：", bold=True)
    para(problem_text)
    para("✅ 解决方案：", bold=True)
    para(solution_text)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. 未来可更新的功能
# ═══════════════════════════════════════════════════════════════
heading("七、未来可更新的功能", level=1)

future = [
    ("转写进度实时反馈",
     "当前转写过程中前端只显示「正在转写中…」，用户无法知道进度。可引入 Server-Sent Events (SSE) 或 WebSocket，将 Whisper 的进度回调实时推送到前端，显示百分比进度条。"),
    ("模型热切换（快速/高精度模式）",
     "预置多个 Whisper 模型尺寸（tiny/base/small/medium），用户可在界面上一键切换「快速模式」与「高精度模式」。快速模式使用 tiny 模型，适合短视频快速出稿。考虑从国内镜像站（ModelScope）下载模型，解决 HuggingFace 网络不稳定问题。"),
    ("GPU 加速支持",
     "当前仅支持 CPU 推理。未来可支持 Apple Silicon（M1/M2/M3）的 MPS 加速或 NVIDIA GPU 的 CUDA 加速，在 Apple Silicon 设备上可获得 3-5 倍转写速度提升。"),
    ("视频预览与片段微调",
     "在选中片段后，用户可在时间轴上拖拽片段的起止点进行微调（如去掉段首的口头禅）。配合波形图可视化，可精确控制剪辑边界。"),
    ("批量上传与队列处理",
     "支持同时上传多个视频文件，后台队列依次处理，用户可切换到任意视频查看结果。适合需要批量处理大量素材的场景（如课程录制、会议记录）。"),
    ("导出格式扩展",
     "除 MP4 和 SRT 外，支持导出 Edius/Premiere 的 EDL（编辑决策列表）、Final Cut Pro XML、DaVinci Resolve 时间线等专业格式，方便在专业软件中二次编辑。"),
    ("背景音乐与转场特效",
     "导出视频时可选择加入背景音乐（自动调整音量避免掩盖人声）、片段间添加平滑转场（淡入淡出、推拉等），提升输出视频的专业观感。"),
    ("字幕样式自定义",
     "在 UI 中提供字幕样式编辑器：字体、大小、颜色、位置、背景色、阴影等，实时预览效果。满足不同平台（抖音、YouTube、B站）的字幕风格需求。"),
    ("关键词自动标签与智能摘要",
     "AI 分析文字稿后自动生成视频内容标签（如#产品介绍 #优惠活动），并生成 2-3 句内容摘要。便于素材归档和快速检索。"),
    ("多用户与云端部署",
     "当前为单机版。未来可容器化部署到服务器（Docker + Nginx），支持多用户同时使用，视频文件存储到云存储（S3/OSS），转写任务分发到 GPU 集群。"),
    ("移动端适配",
     "当前前端仅适配桌面浏览器。可开发移动端响应式布局或 PWA（渐进式 Web 应用），方便用户在手机上快速查看文字稿和管理剪辑任务。"),
]

for title_text, desc in future:
    bullet(f"【{title_text}】{desc}")

# ═══════════════════════════════════════════════════════════════
# 8. 部署说明
# ═══════════════════════════════════════════════════════════════
heading("八、部署与使用说明", level=1)

para("当前软件部署方式：", bold=False)
bullet("位置：/Users/bdh/auto-video-editor/")
bullet("后端：/Users/bdh/auto-video-editor/backend/（FastAPI + uvicorn）")
bullet("前端：/Users/bdh/auto-video-editor/frontend/index.html")
bullet("启动方式：macOS launchd 开机自启（plist 位置：~/Library/LaunchAgents/com.autovideo.server.plist）")
bullet("访问地址：浏览器打开 http://localhost:8000")
bullet("日志位置：/tmp/autovideo.log")
bullet("环境变量：DEEPSEEK_API_KEY + HF_TOKEN（配置在 .env 和 launchd plist 中）")

para("")
para("常用管理命令：", bold=True)
bullet("停止服务：launchctl unload ~/Library/LaunchAgents/com.autovideo.server.plist")
bullet("启动服务：launchctl load ~/Library/LaunchAgents/com.autovideo.server.plist")
bullet("查看日志：tail -f /tmp/autovideo.log")
bullet("检查状态：launchctl list | grep autovideo")

# ═══════════════════════════════════════════════════════════════
# 9. 总结
# ═══════════════════════════════════════════════════════════════
heading("九、总结", level=1)

para("本项目从零到一构建了一套完整的 AI 视频自动剪辑系统。核心创新在于将传统视频剪辑的「时间轴拖拽操作」转变为「自然语言对话」，大幅降低了视频剪辑的使用门槛。")
para("")
para("技术层面，项目集成了语音识别（faster-whisper）、说话人分离（pyannote）、大语言模型（DeepSeek）、视频处理（FFmpeg）等多种 AI 与多媒体技术，并通过 FastAPI 提供统一的 Web 服务接口。部署层面，通过 macOS launchd 实现了开机自启与崩溃自动恢复，用户无需任何技术背景即可使用。")
para("")
para("开发过程中克服了 Python 版本兼容性、依赖冲突、网络不稳定、视频编码精度、CPU 推理性能优化等一系列实际工程问题，积累了完整的端到端交付经验。")
para("")
para("当前版本已具备生产可用性，未来可在实时反馈、性能加速、专业格式导出、云端部署等方向持续演进。")

# ═══════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════
output_path = "/Users/bdh/Desktop/自动视频剪辑工具_项目总结文档.docx"
doc.save(output_path)
print(f"文档已生成：{output_path}")
